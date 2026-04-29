
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r'C:\AfriserveBackend\BUS 1104-01 - AY2026-T4Assignment Activity Unit 2 (Revised).docx'

doc = Document()

# Page setup: US Letter, 1.25 inch L/R margins (like original), 1 inch T/B
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

def set_font(run, bold=False, size=12, font_name='Times New Roman'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_para(doc, text='', bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             indent_first=True, size=12, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(24)  # double spacing
    pf.space_after = Pt(spacing_after)
    pf.space_before = Pt(0)
    if indent_first:
        pf.first_line_indent = Inches(0.5)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size)
    return p

def add_heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(24)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    set_font(run, bold=True)
    return p

def add_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths_inches):
    total_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=total_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Inches(col_widths_inches[i])
        shade_cell(cell, 'D9D9D9')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, bold=True, size=11)
    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.width = Inches(col_widths_inches[ci])
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            run = p.add_run(val)
            set_font(run, size=11)
    return table

# ===================== TITLE PAGE =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.line_spacing = Pt(24)
run = p.add_run('Assignment: Analysis of Unemployment and Inflation Trends in Kenya (2014\u20132024)')
set_font(run, bold=True, size=14)

for line in ['', 'Course Title: Macroeconomics', 'Instructor Name: Shabbir Ahmed', 'Date: 04/20/2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(line)
    set_font(run)

doc.add_page_break()

# ===================== INTRO =====================
add_para(doc, (
    'Over the past ten years, Kenya\u2019s economy has gone through some major changes, particularly in '
    'unemployment and inflation. Both indicators matter a lot\u2014they affect ordinary people\u2019s lives '
    'and reflect the overall health of the economy. This report examines how unemployment is measured in Kenya, '
    'the trends from 2014 to 2024, the types of unemployment in the country, and how inflation affects the cost '
    'of living. It also discusses what these patterns mean for a developing country like Kenya, compared to what '
    'they would look like in a developed one.'
))

# ===================== SECTION 1 =====================
add_heading(doc, '1. Unemployment: Current Rate and Measurement Method')
add_para(doc, (
    'As of early 2024, Kenya\u2019s unemployment rate is estimated at about 12.7% (Kenya National Bureau of Statistics '
    '[KNBS], 2024). Even so, this number does not fully capture the real situation, because many Kenyans work in '
    'informal jobs with low and unstable income. They may technically count as employed, but they are still '
    'struggling financially.'
))
add_para(doc, (
    'The Kenya National Bureau of Statistics (KNBS) measures unemployment through the Labour Force Survey, '
    'following international standards set by the International Labour Organization (ILO). People aged 16 and above '
    'are grouped into three categories: employed, unemployed, and not in the labour force. This method is widely '
    'accepted, but it does not fully account for underemployment\u2014which is common in Kenya, where many workers '
    'take on casual or low-skill jobs that do not match their qualifications (Shapiro et al., 2023).'
))

# ===================== SECTION 2 =====================
add_heading(doc, '2. Unemployment Trends (2014\u20132024) and Influencing Factors')
add_para(doc, 'Table 1 below summarizes Kenya\u2019s estimated unemployment rate over the past decade, along with the key driver for each year.')

doc.add_paragraph()
p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap.paragraph_format.space_after = Pt(4)
run = p_cap.add_run('Table 1: Kenya Unemployment Rate Trends (2014\u20132024)')
set_font(run, bold=True, size=11)

unemp_rows = [
    ('2014', '12.5%', 'Pre-growth stable period'),
    ('2015', '12.4%', 'Continued stability'),
    ('2016', '11.5%', 'Tech sector expansion'),
    ('2017', '11.4%', 'Construction growth'),
    ('2018', '11.2%', 'Digital innovation jobs'),
    ('2019', '10.4%', 'Pre-pandemic low'),
    ('2020', '13.1%', 'COVID-19 impact'),
    ('2021', '12.5%', 'Slow recovery'),
    ('2022', '12.9%', 'High inflation drag'),
    ('2023', '12.6%', 'Ongoing recovery'),
    ('2024', '12.7%', 'Current estimate'),
]
add_table(doc, ['Year', 'Unemployment Rate', 'Key Driver'], unemp_rows, [0.8, 1.6, 3.6])

p_src = doc.add_paragraph()
p_src.paragraph_format.space_before = Pt(4)
p_src.paragraph_format.line_spacing = Pt(24)
run = p_src.add_run('Source: KNBS (2024); World Bank (2024).')
set_font(run, size=10)

doc.add_paragraph()
add_para(doc, (
    'Between 2014 and 2019, unemployment was relatively stable and even declined slightly. Growth in construction '
    'and digital technology, especially Nairobi\u2019s tech scene, created new jobs. But 2020 changed everything. '
    'The COVID-19 pandemic forced many businesses to shut down or scale back, causing unemployment to spike to around '
    '13.1% (Gourinchas, 2023). Recovery from 2022 onwards has been slow, partly because high inflation and rising '
    'taxes have made it harder for businesses to expand and hire.'
))
add_para(doc, (
    'Two structural factors also explain the persistent unemployment. First, Kenya depends heavily on agriculture. '
    'When droughts or poor rains hit, many agricultural workers lose their livelihoods quickly. Second, Kenya has a '
    'rapidly growing youth population, and the economy is simply not creating enough jobs fast enough to absorb them '
    'all. This mismatch leads to high youth unemployment and widespread underemployment (Ha et al., 2019).'
))

# ===================== SECTION 3 =====================
add_heading(doc, '3. Classification of Unemployment: Short Run vs. Long Run')
add_para(doc, (
    'In the short run, Kenya is experiencing cyclical unemployment\u2014unemployment driven by downturns in economic '
    'activity. The COVID-19 pandemic is the clearest example: reduced demand across sectors led businesses to cut '
    'workers. When economic activity bounces back, this type of unemployment typically falls (Gourinchas, 2023).'
))
add_para(doc, (
    'In the long run, structural unemployment is the bigger problem. This happens when there is a mismatch between '
    'the skills workers have and what employers actually need. In Kenya, many graduates lack the practical or '
    'technical skills demanded by sectors like technology and manufacturing. No matter how well the economy performs, '
    'these workers remain unemployed or underemployed unless the skills gap is addressed.'
))
add_para(doc, (
    'The key difference between the two is what causes them. Short-run unemployment is tied to economic cycles: '
    'downturns cause it, recoveries reduce it. Long-run structural unemployment is caused by deeper issues\u2014gaps '
    'in education quality, lack of vocational training, and limited infrastructure in rural areas. These problems '
    'take years, not months, to fix (Shapiro et al., 2023).'
))

# ===================== SECTION 4 =====================
add_heading(doc, '4. Inflation Trends (2014\u20132024): Measurement and Indices')
add_para(doc, 'Table 2 shows Kenya\u2019s estimated annual inflation rate over the same period, alongside the main price driver for each year.')

doc.add_paragraph()
p_cap2 = doc.add_paragraph()
p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap2.paragraph_format.space_after = Pt(4)
run = p_cap2.add_run('Table 2: Kenya Annual Inflation Rate Trends (2014\u20132024)')
set_font(run, bold=True, size=11)

infl_rows = [
    ('2014', '6.9%', 'Food price pressure'),
    ('2015', '6.6%', 'Moderate prices'),
    ('2016', '6.3%', 'Declining trend'),
    ('2017', '8.0%', 'Drought effect on food'),
    ('2018', '4.7%', 'Improved food supply'),
    ('2019', '5.2%', 'Stable period'),
    ('2020', '5.4%', 'Pandemic supply chains'),
    ('2021', '6.1%', 'Post-pandemic pressures'),
    ('2022', '7.7%', 'Energy & food spike'),
    ('2023', '7.7%', 'Elevated cost of living'),
    ('2024', '5.5%', 'Gradual decline'),
]
add_table(doc, ['Year', 'Inflation Rate (CPI)', 'Key Driver'], infl_rows, [0.8, 1.6, 3.6])

p_src2 = doc.add_paragraph()
p_src2.paragraph_format.space_before = Pt(4)
p_src2.paragraph_format.line_spacing = Pt(24)
run = p_src2.add_run('Source: IMF (2023); KNBS (2024).')
set_font(run, size=10)

doc.add_paragraph()
add_para(doc, (
    'Kenya measures inflation using the Consumer Price Index (CPI), which tracks price changes across a standard '
    'basket of goods and services. Food carries about a third of the CPI weight, so food price swings\u2014often '
    'caused by drought or global commodity markets\u2014have a large impact on overall inflation figures.'
))
add_para(doc, (
    'There are two commonly used measures. Headline inflation includes all items in the basket, so it can swing '
    'sharply when food or fuel prices change. Core inflation removes food and energy to reveal the underlying '
    'price trend and is more useful for assessing monetary policy (IMF, 2023). Another useful tool is the Producer '
    'Price Index (PPI), which tracks price changes at the production level before goods reach consumers. Rising PPI '
    'figures can signal future consumer price increases, giving policymakers advance warning.'
))
add_para(doc, (
    'Looking at the table, inflation was moderate from 2014 to 2019, with a notable spike in 2017 due to a severe '
    'drought. After the pandemic disrupted supply chains in 2020 and 2021, inflation climbed sharply, reaching 7.7% '
    'in 2022 and 2023, driven mainly by global energy costs and food prices. By 2024, it had started to ease.'
))

# ===================== SECTION 5 =====================
add_heading(doc, '5. Economic Implications: Kenya vs. a Developed Country')
add_para(doc, (
    'For Kenya, high unemployment and inflation are especially damaging because most workers are in the informal '
    'sector. Their incomes do not automatically adjust when prices rise, so inflation directly erodes their '
    'purchasing power. Household budgets shrink, fewer children stay in school, and overall well-being declines '
    '(Shapiro et al., 2023).'
))
add_para(doc, (
    'High inflation also discourages businesses from investing, since future costs become uncertain. This slows job '
    'creation and makes the unemployment problem worse (Ha et al., 2019). The combination creates a difficult cycle: '
    'fewer jobs mean less consumer spending, and less spending further slows growth.'
))
add_para(doc, (
    'If Kenya were a developed country, these implications would look quite different. Developed economies have '
    'stronger social safety nets\u2014unemployment insurance, food assistance, and healthcare coverage\u2014that '
    'cushion the blow when people lose jobs or when prices rise. Their central banks also have more credibility and '
    'better tools to manage inflation, so inflation tends to be lower and more stable. Structural unemployment still '
    'exists in developed countries, but robust retraining programs and better-funded education systems help workers '
    'adapt faster. In short, the same economic shocks hit developing countries harder and take longer to recover from.'
))

# ===================== CONCLUSION =====================
add_heading(doc, 'Conclusion')
add_para(doc, (
    'Unemployment and inflation remain two of Kenya\u2019s most pressing economic challenges. The official 12.7% '
    'unemployment figure understates the full picture, especially given how widespread informal and precarious work '
    'is. Over the past decade, shocks like the COVID-19 pandemic and rising global commodity prices have made things '
    'harder. In the long run, structural mismatches between education and labor market needs remain the central issue. '
    'Addressing all of this will require sustained investment in education, job creation, and smarter monetary '
    'policy\u2014not quick fixes.'
))

# ===================== REFERENCES =====================
p_ref_hdr = doc.add_paragraph()
p_ref_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ref_hdr.paragraph_format.space_before = Pt(24)
p_ref_hdr.paragraph_format.line_spacing = Pt(24)
run = p_ref_hdr.add_run('References')
set_font(run, bold=True)

refs = [
    'Gourinchas, P.-O. (2023). Global economy on track but not yet out of the woods. IMF Blog. https://www.imf.org/en/Blogs/Articles/2023/04/11/global-economy-is-on-track-but-not-yet-out-of-the-woods',
    'Ha, J., Kose, M. A., & Ohnsorge, F. L. (2019). Understanding inflation in emerging and developing economies. World Bank Group. https://doi.org/10.1596/978-1-4648-1375-7',
    'International Monetary Fund. (2023). Inflation and price stability. IMF. https://www.imf.org/en/About/Factsheets/Sheets/2023/inflation-price-stability',
    'Kenya National Bureau of Statistics. (2024). Economic survey 2024. KNBS. https://www.knbs.or.ke',
    'Shapiro, D., Greenlaw, S., & MacDonald, D. (2023). Principles of macroeconomics (3rd ed.). OpenStax. https://openstax.org/books/principles-macroeconomics-3e/pages/1-introduction',
    'World Bank. (2024). Kenya overview. World Bank Group. https://www.worldbank.org/en/country/kenya/overview',
]
for ref in refs:
    add_ref(doc, ref)

doc.save(OUT)
print('Saved to', OUT)
